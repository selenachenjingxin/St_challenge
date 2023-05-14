import streamlit as st
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
import io

st.title("一键文档优化工具")

# 优化内容的prompt template
template1 = """
你是一位专业的技术文案编辑，现在我希望你对以下文本进行优化：
- 确保语法正确
- 确保拼写正确且一致
- 确保词语使用得当

并输出优化后的内容。

文本如下：{text1}
"""
prompt1=PromptTemplate(template=template1, input_variables=["text1"])

# 易懂+结构化
template2 = """
你是一位专业的技术文案编辑，现在我希望你将以下文本进行结构化并且易于理解。无需解释，直接输出优化后的文本。

文本如下：{text2}
"""
prompt2=PromptTemplate(template=template2, input_variables=["text2"])

uploaded_file = st.file_uploader("上传一个 Word 文件", type=["docx"])

llm = OpenAI(temperature=0.7, max_tokens=2500)
chain1 = LLMChain(llm=llm, prompt=prompt1)
chain2 = LLMChain(llm=llm, prompt=prompt2)
sequential_chain = SimpleSequentialChain(chains=[chain1, chain2])

# Add checkboxes for options
spell_check = st.checkbox("拼写检查")
grammar_check = st.checkbox("语法检查")
punctuation_check = st.checkbox("标点符号检查")
sentence_structure_optimization = st.checkbox("句子结构优化")
paragraph_structure_check = st.checkbox("段落结构检查")

show_results = st.checkbox("显示处理结果", value=True)
process_file_button = st.button("开始处理文件")

if process_file_button and uploaded_file is not None:
    document = docx.Document(uploaded_file)
    new_document = docx.Document()

    for para in document.paragraphs:
        # Modify the prompt template based on the selected options
        prompt_text = ""
        if spell_check:
            prompt_text += "- 进行拼写检查和纠正\n"
        if grammar_check:
            prompt_text += "- 进行语法检查和修正\n"
        if punctuation_check:
            prompt_text += "- 进行标点符号检查和优化\n"
        if sentence_structure_optimization:
            prompt_text += "- 进行句子结构优化\n"
        if paragraph_structure_check:
            prompt_text += "- 进行段落结构检查和优化\n"

        # Update the prompt template with the selected options and run the sequential chain
        prompt1.update_template({"text1": para.text, "options": prompt_text})
        response = sequential_chain.run(prompt=prompt1)

        if show_results:
            st.write("原文本：")
            st.write(para.text)
            st.write("优化后的文本：")
            st.write(response)

        new_document.add_paragraph(response)

    new_docx = io.BytesIO()
    new_document.save(new_docx)
    new_docx.seek(0)

    st.download_button(
        label="下载优化后的 Word 文件",
        data=new_docx,
        file_name='optimized_document.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
