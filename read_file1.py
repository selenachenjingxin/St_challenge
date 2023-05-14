import streamlit as st
import io
import base64
from docx import Document
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain, SimpleSequentialChain

# 优化内容的prompt template
template1 = """
you are a professional technical writer, You are now tasked with optimizing the text. This involves checking and correcting the content to ensure:

Correct grammar.
Consistent and accurate spelling.
Appropriate use of words.

please do the task based on text:{text1}.
"""

prompt1 = PromptTemplate(template=template1, input_variables=["text1"])

template2 = """
As a professional technical writer, please make the text more structured and easy to understand.
Just provide the revised text in Chinese with no explanation.

please do the task based on text:{text2}.
"""

prompt2 = PromptTemplate(template=template2, input_variables=["text2"])

llm = OpenAI(temperature=0.7, max_tokens=2500)
chain1 = LLMChain(llm=llm, prompt=prompt1)
chain2 = LLMChain(llm=llm, prompt=prompt2)
sequential_chain = SimpleSequentialChain(chains=[chain1, chain2])

#转化为HTML
template3 = """
You are a professional technical writer, now I want you to convert the text into a HTML style.

please do the task based on text:{text3}.
"""
prompt3=PromptTemplate(
    template=template3,
    input_variables=["text3"])


def process_paragraphs(paragraph_texts):
    processed_paragraphs = []
    for para_text in paragraph_texts:
        response2 = chain2.run(para_text)
        processed_paragraphs.append(response2)
    return processed_paragraphs


def download_link(object_to_download, download_filename, download_link_text):
    b64 = base64.b64encode(object_to_download).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{download_filename}">{download_link_text}</a>'


st.title('Word 文档优化处理')

uploaded_file = st.file_uploader("选择一个 Word 文档", type=['docx'])

if uploaded_file is not None:
    document = Document(uploaded_file)
    paragraphs = document.paragraphs
    paragraph_texts = [p.text for p in paragraphs]
    st.write('原始文档内容：')
    for p in paragraph_texts:
        st.write(p)

    processed_paragraphs = process_paragraphs(paragraph_texts)

    st.write('处理后的文档内容：')
    for p in processed_paragraphs:
        st.write(p)

    if st.button('生成处理后的 Word 文档'):
        new_doc = Document()
        for p in processed_paragraphs:
            new_doc.add_paragraph(p)

        output = io.BytesIO()
        new_doc.save(output)

        download_link_html = download_link(output.getvalue(), 'processed_document.docx', '点击这里下载处理后的 Word 文档')
        st.markdown(download_link_html, unsafe_allow_html=True)
