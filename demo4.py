import streamlit as st
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
import io

st.title("一键文档优化工具")

# 定义各种优化功能的prompt template
templates = {
    "拼写检查优化": "你是一位专业的技术文档工程师，现在我希望你对以下文本进行拼写检查，修正不正确的用词用字和句子，并输出修正后的内容。文本如下：{text}",
    "语法检查优化": "你是一位专业的技术文档工程师，现在我希望你对以下文本进行语法检查，修正不正确的语法，并输出优化后的内容。文本如下：{text}",
    "标点符号检查": "你是一位专业的技术文档工程师，现在我希望你对以下文本进行标点符号检查，修改使用不恰当的标点符号，并输出优化后的内容。文本如下：{text}",
    "句子结构优化": "你是一位专业的技术文档工程师，现在我希望你对以下文本进行句子结构优化，若有长难句，简化长难句并输出优化后的内容。文本如下：{text}",
    "段落结构优化": "你是一位专业的技术文档工程师，现在我希望你对以下文本进行段落结构优化，将复杂的段落，在可能的情况下使用有序列表或者无序列表等，并输出优化后的内容。文本如下：{text}",
    "自动摘要": "你是一位专业的技术文档工程师，现在我希望你根据以下文本，生成一段简短的摘要。文本如下：{text}",
    "文档格式优化": "你是一位专业的技术文档工程师，现在我希望你根据以下文本，给出适当的格式化建议。文本如下：{text}",
    "内容翻译": "你是一位专业的技术文档工程师，现在我希望你将以下文本翻译成{language}。文本如下：{text}",
}

# 创建复选框供用户选择需要的优化功能
options = st.multiselect("选择需要的优化功能", list(templates.keys()))

# 添加语言选择的下拉列表
languages = ["英语", "法语", "德语", "西班牙语", "日语", "韩语", "俄语", "阿拉伯语", "葡萄牙语", "意大利语"]
selected_language = st.selectbox("选择翻译的目标语言（如果需要的话）", languages)

llm = OpenAI(temperature=0.7, max_tokens=2500)
chains = []

# 根据用户选择的优化功能生成对应的模型链
for option in options:
    template = templates[option]
    input_variables = ["text"]
    if option == "内容翻译":
        # 第一步：翻译
        translation_template = "将以下文本翻译成{language}：{text}"
        translation_prompt = PromptTemplate(template=translation_template, input_variables=input_variables)
        translation_chain = LLMChain(llm=llm, prompt=translation_prompt)
        chains.append(translation_chain)

        # 第二步：优化
        option = "文本优化"  # 假设我们有一个名为"文本优化"的模板用于优化翻译后的文本
        template = templates[option]
    prompt = PromptTemplate(template=template, input_variables=input_variables)
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

sequential_chain = SimpleSequentialChain(chains=chains)

# Add a text area for user to paste the text
input_text = st.text_area("在此处粘贴文本进行优化", "")

# If there is any input text, process it
if input_text:
    response = sequential_chain.run(input_text)
    st.write("优化后的文本：")
    st.markdown(response)

# Add a file uploader for user to upload a docx file
uploaded_file = st.file_uploader("或者上传一个 Word 文件", type=["docx"])

# If there is any uploaded file, process it
if uploaded_file is not None:
    document = docx.Document(uploaded_file)
    new_document = docx.Document()

    for para in document.paragraphs:
        response = sequential_chain.run(para.text)
        new_document.add_paragraph(response)

    new_docx = io.BytesIO()
    new_document.save(new_docx)
    new_docx.seek(0)

    st.download_button(
        label="下载优化后的 Word 文件",
        data=new_docx,
        file_name='optimized_document.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
