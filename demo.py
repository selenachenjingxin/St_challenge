import streamlit as st
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
import io

st.title("一键文档优化工具")

# 定义各种优化功能的prompt template
templates = {
    "拼写检查": "你是一位专业的技术文案编辑，现在我希望你对以下文本进行拼写检查，修正不争取的用词用字，并输出优化后的内容。文本如下：{text}",
    "语法检查": "你是一位专业的技术文案编辑，现在我希望你对以下文本进行语法检查，修正不正确的语法，并输出优化后的内容。文本如下：{text}",
    "标点符号检查": "你是一位专业的技术文案编辑，现在我希望你对以下文本进行标点符号检查，修改使用不恰当的标点符号，并输出优化后的内容。文本如下：{text}",
    "句子结构优化": "你是一位专业的技术文案编辑，现在我希望你对以下文本进行句子结构优化，若有长难句，简化长难句并输出优化后的内容。文本如下：{text}",
    "段落结构检查": "你是一位专业的技术文案编辑，现在我希望你对以下文本进行段落结构检查，将复杂的段落，在可能的情况下使用有序列表或者无序列表等，并输出优化后的内容。文本如下：{text}",
}

# 创建复选框供用户选择需要的优化功能
options = st.multiselect("选择需要的优化功能", list(templates.keys()))

uploaded_file = st.file_uploader("上传一个 Word 文件", type=["docx"])

llm = OpenAI(temperature=0.7, max_tokens=2500)
chains = []

# 根据用户选择的优化功能生成对应的模型链
for option in options:
    template = templates[option]
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

sequential_chain = SimpleSequentialChain(chains=chains)

show_results = st.checkbox("显示处理结果", value=True)
process_file_button = st.button("开始处理文件")

if process_file_button and uploaded_file is not None:
    document = docx.Document(uploaded_file)
    new_document = docx.Document()

    for para in document.paragraphs:
        response = sequential_chain.run(para.text)

        if show_results:
            st.write("原文本：")
            st.write(para.text)
            st.write("优化后的文本：")
            st.write(response)

        new_document.add_paragraph(response)

    new_docx = io.BytesIO()
    new_document.save(new_docx)
    new_docx.seek(0)

    st.download_button(
        label="下载优化后的 Word 文件",
        data=new_docx,
        file_name='optimized_document.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
