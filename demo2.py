import streamlit as st
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
import io

from yaml import DocumentEndEvent

st.title("一键文档优化工具")

llm = OpenAI(temperature=0.7, max_tokens=2500)
chains = []

# 为每一个优化功能创建一个复选框，并根据用户的选择来创建对应的模型链
if st.sidebar.checkbox("拼写检查优化"):
    template = "你是一位专业的技术文档工程师，现在我希望你对以下文本进行拼写检查，修正不正确的用词用字和句子，并输出修正后的内容。文本如下：{text}"
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

if st.sidebar.checkbox("语法检查优化"):
    template = "你是一位专业的技术文档工程师，现在我希望你对以下文本进行语法检查，修正不正确的语法，并输出优化后的内容。文本如下：{text}"
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

if st.sidebar.checkbox("标点符号检查"):
    template = "你是一位专业的技术文档工程师，现在我希望你对以下文本进行标点符号检查，修改使用不恰当的标点符号，并输出优化后的内容。文本如下：{text}"
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

if st.sidebar.checkbox("句子结构优化"):
    template = "你是一位专业的技术文档工程师，现在我希望你对以下文本进行句子结构优化，若有长难句，简化长难句并输出优化后的内容。文本如下：{text}"
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

if st.sidebar.checkbox("段落结构优化"):
    template = "你是一位专业的技术文档工程师，现在我希望你对以下文本进行段落结构优化，将复杂的段落，在可能的情况下使用有序列表或者无序列表等，并输出优化后的内容。文本如下：{text}"
    prompt = PromptTemplate(template=template, input_variables=["text"])
    chain = LLMChain(llm=llm, prompt=prompt)
    chains.append(chain)

sequential_chain = SimpleSequentialChain(chains=chains)


# Add a radio button for navigation
page = st.sidebar.radio("导航", ["优化文本", "优化 Word 文件"])

if page == "优化文本":
    # Add a text area for user to paste the text
    input_text = st.text_area("在此处粘贴文本进行优化", "", height=300)

    # Create two columns for displaying the original and optimized text side by side
    col1, col2 = st.columns(2)

    # Add a button to trigger the processing
    if st.button('优化文本'):
        if input_text:
            response = sequential_chain.run(input_text)
            col1.write("原始文本：")
            col1.markdown(input_text)
            col2.write("优化后的文本：")
            col2.markdown(response)

elif page == "优化 Word 文件":
    # Add a file uploader for user to upload a docx file
    uploaded_file = st.file_uploader("上传一个 Word 文件", type=["docx"])

    # If there is any uploaded file, process it
    if uploaded_file is not None:
        document = docx.Document(uploaded_file)
        # Read the document and concatenate all the paragraphs into a single string
        document_text = "\n".join([para.text for para in document.paragraphs])

        # Run the sequential chain on the concatenated text
        response = sequential_chain.run(document_text)

        # Create a new document and add the optimized text as a single paragraph
        new_document = docx.Document()
        new_document.add_paragraph(response)

        # Save the optimized document
        new_docx = io.BytesIO()
        new_document.save(new_docx)
        new_docx.seek(0)

        # Download the optimized document
        if st.button('下载优化后的 Word 文件'):
            st.download_button(
                label="点击下载",
                data=new_docx,
                file_name='optimized_document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
