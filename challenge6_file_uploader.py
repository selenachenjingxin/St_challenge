import streamlit as st
import pandas as pd
from io import StringIO
import io
import docx
import os
import openai
from langchain.llms import OpenAI
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.llms import OpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.document_loaders import UnstructuredHTMLLoader
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain,SequentialChain
from langchain.memory import ConversationBufferMemory
from langchain.chains import SimpleSequentialChain

st.title("一键文档优化工具")

## 优化内容的prompt template
template1 = """
You are a professional technical writer, now I want you to optimize text which means to check and correct the content to make sure that:

 Grammar is correct.
 Spelling is correct and consistent.
 word is rightly used.
 
and just output the optimized content.

please do the task based on text:{text1}.
"""
prompt1=PromptTemplate(
    template=template1,
    input_variables=["text1"])

# 易懂+结构化
template4 = """
You are a professional technical writer, now I want you to make the text more structured and easy to understand.
Just output the revised text in Chinese with no explanation.

please do the task based on text:{text4}.
"""
prompt4=PromptTemplate(
    template=template4,
    input_variables=["text4"])

# 添加一个用于上传文件的 Streamlit 小部件
uploaded_file = st.file_uploader("上传一个 Word 文件", type=["docx"])

        
from langchain.text_splitter import CharacterTextSplitter
text_splitter = CharacterTextSplitter(        
    separator = "\n\n",
    chunk_size = 1000,
    chunk_overlap  = 200,
    length_function = len,
)

llm = OpenAI(temperature=0.7,max_tokens=2500)
chain1 = LLMChain(llm=llm, prompt=prompt1)
chain4 = LLMChain(llm=llm, prompt=prompt4)
sequential_chain=SimpleSequentialChain(chains=[chain1,chain4])

# Add a button for the user to click after uploading a file
process_file_button = st.button("开始处理文件")

# 读取文件
if process_file_button:
    document = docx.Document(uploaded_file)

    # 创建一个新的 Word 文档来保存处理后的段落
    new_document = docx.Document()

    # 遍历原始文档中的每个段落
    for para in document.paragraphs:
        st.write(para.text)

        # 处理段落
        response2 = sequential_chain.run(para.text)
        st.write(response2)

        # 将处理后的段落添加到新文档中
        new_document.add_paragraph(response2)
 


