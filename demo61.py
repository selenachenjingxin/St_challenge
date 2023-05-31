import streamlit as st
import docx
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate
from docx.shared import Pt
import io

st.title("一键文档优化工具")

llm = OpenAI(temperature=0.7, max_tokens=2500)
chains = []

# Add a radio button for navigation
st.sidebar.markdown("<h2>导航</h2>", unsafe_allow_html=True)
page = st.sidebar.radio("", ["优化文本", "优化 Word 文件"])
# 主模板+附加模板
main_template = """
你是一位专业的技术文档工程师，现在我希望你对以下文本进行优化。文本如下：
{text}
请按照以下优化要求进行优化：
1. 拼写检查：修正错别字和不正确的用词。
2. 语法检查：修正不正确的语法。
3. 标点符号检查：修改使用不恰当的标点符号。
4. 句子结构优化：简化长难句。
5. 段落结构优化：使用有序列表或无序列表等优化段落结构。
6. 为文本撰写合适的标题
请输出优化后的文本。
"""
prompt = PromptTemplate(template=main_template, input_variables=["text"])
chain = LLMChain(llm=llm, prompt=prompt)
chains.append(chain)

st.sidebar.header("优化功能选择：")
st.sidebar.markdown("### 内容质量优化（默认）")

st.sidebar.checkbox("拼写检查", value=True)
st.sidebar.checkbox("语法检查",value=True)
st.sidebar.checkbox("标点符号检查",value=True)
st.sidebar.checkbox("句子结构优化",value=True)
st.sidebar.checkbox("段落结构优化",value=True)
st.sidebar.checkbox("术语检查", value=False)
st.sidebar.checkbox("语义分析", value=False)
st.sidebar.checkbox("文本去重", value=False)


sequential_chain = SimpleSequentialChain(chains=chains)


if page == "优化文本":
    # Add a text area for user to paste the text
    input_text = st.text_area("在此处粘贴文本进行优化", "", height=300)

    # Create two columns for displaying the original and optimized text side by side
    col1, col2 = st.columns(2)

    # Add a button to trigger the processing
    if st.button('优化文本'):
        if input_text:
            response = sequential_chain.run(input_text)
            col1.write("原始文本：")
            col1.markdown(input_text)
            col2.write("优化后的文本：")
            col2.markdown(response)

else:
    # Add a file uploader for user to upload a docx file
    uploaded_file = st.file_uploader("上传一个 Word 文件", type=["docx"])

        # If there is any uploaded file, process it
    if uploaded_file is not None:
        document = docx.Document(uploaded_file)

        # Read the document and concatenate all the paragraphs into a single string
        document_text = "\n".join([para.text for para in document.paragraphs])

        # Run the sequential chain on the concatenated text
        response = sequential_chain.run(document_text)

        # Create a new document and add the optimized text as a single paragraph
        new_document = docx.Document()
        
        # Set document font size and style
        run = new_document.add_paragraph().add_run()
        font = run.font
        font.size = Pt(12)
        font.name = 'Arial'

        # Add formatted paragraphs to the new document
        paragraphs = response.split("\n")
        for paragraph in paragraphs:
            new_document.add_paragraph(paragraph)

        # Save the optimized document
        new_docx = io.BytesIO()
        new_document.save(new_docx)
        new_docx.seek(0)

        # Create an empty element for the download button
        download_button_placeholder = st.empty()

        # Update the placeholder with the actual download button
        if download_button_placeholder.button('下载优化后的 Word 文件'):
            st.download_button(
                label="点击下载",
                data=new_docx,
                file_name='optimized_document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
